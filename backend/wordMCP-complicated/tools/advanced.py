"""
Advanced tools for Word documents.
"""

from typing import Optional, List
import logging
import re

from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Import from core - handle both package and direct execution
try:
    from ..core.path_utils import PathUtils
    from ..core.exceptions import DocumentError, DocumentOperationError
    from ..core.document import DocumentManager
except ImportError:
    from core.path_utils import PathUtils
    from core.exceptions import DocumentError, DocumentOperationError
    from core.document import DocumentManager

logger = logging.getLogger(__name__)


def register_advanced_tools(mcp: FastMCP) -> None:
    """
    Register advanced tools with the MCP server.
    
    Args:
        mcp: FastMCP server instance
    """
    path_utils = PathUtils()
    doc_manager = DocumentManager()
    
    @mcp.tool()
    def insert_image(
        file_path: str,
        image_path: str,
        width: Optional[float] = None,
        caption: Optional[str] = None
    ) -> dict:
        """
        Insert an image into a Word document.
        
        Args:
            file_path: Path to the Word document
            image_path: Path to the image file (jpg, jpeg, png, gif, bmp)
            width: Width of image in inches (maintains aspect ratio)
            caption: Optional caption text below the image
        
        Returns:
            Dictionary with operation result
        
        Examples:
            insert_image("report.docx", "/path/to/image.png", width=5.0)
            insert_image("report.docx", "logo.jpg", width=3.0, caption="Company Logo")
        """
        try:
            logger.info(f"Tool called: insert_image(file_path={file_path}, image={image_path})")
            return doc_manager.insert_image(file_path, image_path, width, caption)
        except DocumentError as e:
            logger.error(f"Document error in insert_image: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in insert_image: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def search_text(
        file_path: str,
        search_text: str,
        match_case: bool = False
    ) -> dict:
        """
        Search for text in a Word document.
        
        Args:
            file_path: Path to the Word document
            search_text: Text to search for
            match_case: Whether to match case (default: False)
        
        Returns:
            Dictionary with search results including paragraph indices and text
        
        Examples:
            search_text("report.docx", "summary")
            search_text("report.docx", "Important", match_case=True)
        """
        try:
            logger.info(f"Tool called: search_text(file_path={file_path}, search={search_text})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Search for text
            matches = []
            pattern = re.compile(
                re.escape(search_text),
                re.IGNORECASE if not match_case else 0
            )
            
            for i, para in enumerate(doc.paragraphs):
                if para.text:
                    found = pattern.findall(para.text)
                    if found:
                        matches.append({
                            "paragraph_index": i,
                            "text": para.text,
                            "match_count": len(found)
                        })
            
            result = {
                "success": True,
                "file_path": str(doc_path),
                "search_text": search_text,
                "match_case": match_case,
                "total_matches": sum(m["match_count"] for m in matches),
                "paragraphs_with_matches": len(matches),
                "matches": matches
            }
            
            logger.info(f"Search completed: {len(matches)} paragraphs with matches")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in search_text: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in search_text: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def replace_text(
        file_path: str,
        search_text: str,
        replace_text: str,
        match_case: bool = False,
        max_replacements: Optional[int] = None
    ) -> dict:
        """
        Search and replace text in a Word document.
        
        Args:
            file_path: Path to the Word document
            search_text: Text to search for
            replace_text: Text to replace with
            match_case: Whether to match case (default: False)
            max_replacements: Maximum number of replacements (None = unlimited)
        
        Returns:
            Dictionary with operation result and replacement count
        
        Examples:
            replace_text("report.docx", "old text", "new text")
            replace_text("report.docx", "ABC", "XYZ", match_case=True, max_replacements=5)
        """
        try:
            logger.info(f"Tool called: replace_text(file_path={file_path})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Replace text
            replacement_count = 0
            flags = 0 if match_case else re.IGNORECASE
            
            for para in doc.paragraphs:
                for run in para.runs:
                    if search_text in run.text or (not match_case and search_text.lower() in run.text.lower()):
                        if max_replacements and replacement_count >= max_replacements:
                            break
                        
                        # Count replacements before doing them
                        matches = len(re.findall(re.escape(search_text), run.text, flags))
                        
                        # Perform replacement
                        run.text = re.sub(
                            re.escape(search_text),
                            replace_text,
                            run.text,
                            count=max_replacements - replacement_count if max_replacements else 0,
                            flags=flags
                        )
                        
                        replacement_count += matches
                        
                        if max_replacements and replacement_count >= max_replacements:
                            break
                
                if max_replacements and replacement_count >= max_replacements:
                    break
            
            # Save document
            doc.save(str(doc_path))
            
            result = {
                "success": True,
                "message": f"Replaced {replacement_count} occurrence(s)",
                "file_path": str(doc_path),
                "search_text": search_text,
                "replace_text": replace_text,
                "replacement_count": replacement_count
            }
            
            logger.info(f"Text replaced: {replacement_count} occurrences")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in replace_text: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in replace_text: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def merge_documents(
        output_path: str,
        file_paths: List[str],
        add_page_breaks: bool = True
    ) -> dict:
        """
        Merge multiple Word documents into one.
        
        Args:
            output_path: Path for the merged output document
            file_paths: List of document paths to merge (in order)
            add_page_breaks: Add page breaks between merged documents (default: True)
        
        Returns:
            Dictionary with operation result
        
        Examples:
            merge_documents("combined.docx", ["doc1.docx", "doc2.docx", "doc3.docx"])
            merge_documents("full_report.docx", ["intro.docx", "body.docx", "conclusion.docx"], add_page_breaks=False)
        """
        try:
            logger.info(f"Tool called: merge_documents(output={output_path}, count={len(file_paths)})")
            
            if not file_paths:
                raise DocumentOperationError(
                    "merge_documents",
                    "At least one input file required"
                )
            
            # Create new document
            merged_doc = Document()
            
            # Merge each document
            for i, file_path in enumerate(file_paths):
                try:
                    # Resolve file path
                    doc_path = path_utils.resolve_file_path(file_path)
                    doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
                    
                    # Open document
                    doc = Document(str(doc_path))
                    
                    # Copy paragraphs
                    for para in doc.paragraphs:
                        new_para = merged_doc.add_paragraph(para.text)
                        # Copy paragraph style
                        new_para.style = para.style
                    
                    # Copy tables
                    for table in doc.tables:
                        new_table = merged_doc.add_table(
                            rows=len(table.rows),
                            cols=len(table.columns)
                        )
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                new_table.rows[row_idx].cells[col_idx].text = cell.text
                    
                    # Add page break between documents (except after last)
                    if add_page_breaks and i < len(file_paths) - 1:
                        merged_doc.add_page_break()
                    
                    logger.debug(f"Merged document: {doc_path}")
                    
                except Exception as e:
                    logger.warning(f"Failed to merge {file_path}: {e}")
                    raise DocumentOperationError(
                        "merge_documents",
                        f"Failed to merge {file_path}: {e}"
                    )
            
            # Normalize output path
            output = path_utils.normalize_file_path(output_path)
            output = path_utils.validate_file_path(output, must_exist=False, check_size=False)
            path_utils.ensure_parent_directory(output)
            
            # Save merged document
            merged_doc.save(str(output))
            
            result = {
                "success": True,
                "message": f"Successfully merged {len(file_paths)} documents",
                "output_path": str(output),
                "merged_count": len(file_paths),
                "output_size": output.stat().st_size
            }
            
            logger.info(f"Documents merged successfully: {output}")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in merge_documents: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in merge_documents: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def get_document_stats(file_path: str) -> dict:
        """
        Get detailed statistics about a Word document.
        
        Args:
            file_path: Path to the Word document
        
        Returns:
            Dictionary with document statistics (word count, paragraph count, etc.)
        
        Examples:
            get_document_stats("report.docx")
        """
        try:
            logger.info(f"Tool called: get_document_stats(file_path={file_path})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Calculate statistics
            total_paragraphs = len(doc.paragraphs)
            non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())
            total_text = " ".join(p.text for p in doc.paragraphs)
            word_count = len(total_text.split())
            char_count = len(total_text)
            char_count_no_spaces = len(total_text.replace(" ", ""))
            
            # Count headings
            heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))
            
            # Count tables
            table_count = len(doc.tables)
            total_table_cells = sum(
                len(table.rows) * len(table.columns)
                for table in doc.tables
            )
            
            result = {
                "success": True,
                "file_path": str(doc_path),
                "file_size": doc_path.stat().st_size,
                "statistics": {
                    "word_count": word_count,
                    "character_count": char_count,
                    "character_count_no_spaces": char_count_no_spaces,
                    "paragraph_count": total_paragraphs,
                    "non_empty_paragraphs": non_empty_paragraphs,
                    "heading_count": heading_count,
                    "table_count": table_count,
                    "table_cells": total_table_cells
                }
            }
            
            logger.info(f"Document stats calculated: {word_count} words, {total_paragraphs} paragraphs")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in get_document_stats: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in get_document_stats: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    logger.info("Advanced tools registered successfully")

