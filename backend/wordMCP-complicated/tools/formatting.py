"""
Formatting tools for Word documents.
"""

from typing import Optional
import logging

from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import from core - handle both package and direct execution
try:
    from ..core.path_utils import PathUtils
    from ..core.exceptions import DocumentError, DocumentOperationError
except ImportError:
    from core.path_utils import PathUtils
    from core.exceptions import DocumentError, DocumentOperationError

logger = logging.getLogger(__name__)


def register_formatting_tools(mcp: FastMCP) -> None:
    """
    Register formatting tools with the MCP server.
    
    Args:
        mcp: FastMCP server instance
    """
    path_utils = PathUtils()
    
    @mcp.tool()
    def format_paragraph(
        file_path: str,
        paragraph_index: int,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        color: Optional[str] = None,
        alignment: Optional[str] = None
    ) -> dict:
        """
        Format text in a specific paragraph.
        
        Args:
            file_path: Path to the Word document
            paragraph_index: Index of paragraph to format (0-based)
            font_name: Font name (e.g., "Arial", "Times New Roman")
            font_size: Font size in points
            bold: Make text bold
            italic: Make text italic
            underline: Underline text
            color: Text color in hex format (e.g., "FF0000" for red)
            alignment: Text alignment ("left", "center", "right", "justify")
        
        Returns:
            Dictionary with operation result
        
        Examples:
            format_paragraph("report.docx", 0, font_size=14, bold=True, alignment="center")
            format_paragraph("report.docx", 5, font_name="Arial", color="0000FF", italic=True)
        """
        try:
            logger.info(f"Tool called: format_paragraph(file_path={file_path}, index={paragraph_index})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Check paragraph index
            if paragraph_index >= len(doc.paragraphs):
                raise DocumentOperationError(
                    "format_paragraph",
                    f"Paragraph index {paragraph_index} out of range (max: {len(doc.paragraphs) - 1})",
                    str(doc_path)
                )
            
            para = doc.paragraphs[paragraph_index]
            
            # Apply formatting to all runs in paragraph
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold:
                    run.font.bold = True
                if italic:
                    run.font.italic = True
                if underline:
                    run.font.underline = True
                if color:
                    # Parse hex color
                    try:
                        rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                        run.font.color.rgb = RGBColor(*rgb)
                    except Exception as e:
                        logger.warning(f"Invalid color format: {color}, {e}")
            
            # Apply paragraph alignment
            if alignment:
                alignment_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.lower() in alignment_map:
                    para.alignment = alignment_map[alignment.lower()]
            
            # Save document
            doc.save(str(doc_path))
            
            result = {
                "success": True,
                "message": "Paragraph formatted successfully",
                "file_path": str(doc_path),
                "paragraph_index": paragraph_index
            }
            
            logger.info(f"Paragraph formatted successfully: {doc_path}")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in format_paragraph: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in format_paragraph: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def insert_page_break(file_path: str) -> dict:
        """
        Insert a page break at the end of the document.
        
        Args:
            file_path: Path to the Word document
        
        Returns:
            Dictionary with operation result
        
        Examples:
            insert_page_break("report.docx")
        """
        try:
            logger.info(f"Tool called: insert_page_break(file_path={file_path})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Add page break
            doc.add_page_break()
            
            # Save document
            doc.save(str(doc_path))
            
            result = {
                "success": True,
                "message": "Page break inserted successfully",
                "file_path": str(doc_path)
            }
            
            logger.info(f"Page break inserted: {doc_path}")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in insert_page_break: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in insert_page_break: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def add_bullet_list(
        file_path: str,
        items: list[str],
        title: Optional[str] = None
    ) -> dict:
        """
        Add a bullet list to the document.
        
        Args:
            file_path: Path to the Word document
            items: List of items to add
            title: Optional title before the list
        
        Returns:
            Dictionary with operation result
        
        Examples:
            add_bullet_list("report.docx", ["Item 1", "Item 2", "Item 3"], title="TODO List")
        """
        try:
            logger.info(f"Tool called: add_bullet_list(file_path={file_path})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Add title if provided
            if title:
                doc.add_heading(title, level=3)
            
            # Add bullet items
            for item in items:
                doc.add_paragraph(item, style='List Bullet')
            
            # Save document
            doc.save(str(doc_path))
            
            result = {
                "success": True,
                "message": f"Bullet list with {len(items)} items added successfully",
                "file_path": str(doc_path),
                "item_count": len(items)
            }
            
            logger.info(f"Bullet list added: {doc_path}")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in add_bullet_list: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in add_bullet_list: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    @mcp.tool()
    def add_numbered_list(
        file_path: str,
        items: list[str],
        title: Optional[str] = None
    ) -> dict:
        """
        Add a numbered list to the document.
        
        Args:
            file_path: Path to the Word document
            items: List of items to add
            title: Optional title before the list
        
        Returns:
            Dictionary with operation result
        
        Examples:
            add_numbered_list("report.docx", ["Step 1", "Step 2", "Step 3"], title="Process")
        """
        try:
            logger.info(f"Tool called: add_numbered_list(file_path={file_path})")
            
            # Resolve file path
            doc_path = path_utils.resolve_file_path(file_path)
            doc_path = path_utils.validate_file_path(doc_path, must_exist=True)
            
            # Open document
            doc = Document(str(doc_path))
            
            # Add title if provided
            if title:
                doc.add_heading(title, level=3)
            
            # Add numbered items
            for item in items:
                doc.add_paragraph(item, style='List Number')
            
            # Save document
            doc.save(str(doc_path))
            
            result = {
                "success": True,
                "message": f"Numbered list with {len(items)} items added successfully",
                "file_path": str(doc_path),
                "item_count": len(items)
            }
            
            logger.info(f"Numbered list added: {doc_path}")
            return result
            
        except DocumentError as e:
            logger.error(f"Document error in add_numbered_list: {e}")
            return e.to_dict()
        except Exception as e:
            logger.error(f"Unexpected error in add_numbered_list: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "error_code": "UNEXPECTED_ERROR"
            }
    
    logger.info("Formatting tools registered successfully")

