"""
Word Document MCP Server

This MCP server provides CRUD (Create, Read, Update, Delete) operations
for Microsoft Word documents (.docx files).
"""

from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from pathlib import Path
from typing import Optional, List, Dict
from datetime import datetime

# Initialize FastMCP server
mcp = FastMCP("Word Document MCP Server")


# ==================== Helper Functions ====================

def get_default_word_dir() -> str:
    """Get the default word documents directory (current_dir/word)."""
    current_dir = os.getcwd()
    word_dir = os.path.join(current_dir, "word")
    os.makedirs(word_dir, exist_ok=True)
    return word_dir


def normalize_file_path(file_path: str, default_dir: Optional[str] = None) -> str:
    """
    Normalize file path. If path is relative or just filename, 
    place it in the default word directory.
    
    Args:
        file_path: File path (can be absolute, relative, or just filename)
        default_dir: Default directory to use (defaults to word subdirectory)
    
    Returns:
        Normalized absolute file path
    """
    if default_dir is None:
        default_dir = get_default_word_dir()
    
    # If it's already an absolute path, use it as is
    if os.path.isabs(file_path):
        return file_path
    
    # If it's a relative path with directory, join with default_dir
    if os.path.dirname(file_path):
        return os.path.join(default_dir, file_path)
    
    # If it's just a filename, place it in default_dir
    return os.path.join(default_dir, file_path)


# ==================== Tools ====================

@mcp.tool()
def create_word_document(
    file_path: Optional[str] = None,
    title: Optional[str] = None,
    content: Optional[str] = None
) -> dict:
    """
    Create a new Word document (.docx file).
    
    Args:
        file_path: File path or filename. If not provided or relative, saves to word/ subdirectory.
                   If just filename provided, saves to word/filename.docx
        title: Optional title for the document
        content: Optional initial content for the document
    
    Returns:
        Dictionary with operation result and file path
    """
    try:
        # Use default filename if not provided
        if not file_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_path = f"document_{timestamp}.docx"
        
        # Normalize path (defaults to word/ subdirectory)
        file_path = normalize_file_path(file_path)
        
        # Ensure file path ends with .docx
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(file_path) if os.path.dirname(file_path) else '.', exist_ok=True)
        
        # Create new document
        doc = Document()
        
        # Add title if provided
        if title:
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content if provided
        if content:
            # Split content by newlines and add as paragraphs
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()  # Empty paragraph for blank lines
        
        # Save document
        doc.save(file_path)
        
        return {
            "success": True,
            "message": f"Word document created successfully",
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "created_at": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


@mcp.tool()
def read_word_document(file_path: str) -> dict:
    """
    Read content from a Word document.
    
    Args:
        file_path: Path to the Word document (.docx file). 
                   If relative or just filename, searches in word/ subdirectory first.
    
    Returns:
        Dictionary containing document content, paragraphs, and metadata
    """
    try:
        # Try to normalize path (searches in word/ subdirectory if relative)
        normalized_path = normalize_file_path(file_path)
        
        # Check both normalized path and original path
        if os.path.exists(normalized_path):
            file_path = normalized_path
        elif not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "file_path": file_path
            }
        
        # Open document
        doc = Document(file_path)
        
        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Extract tables
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        # Get document properties
        core_props = doc.core_properties
        
        return {
            "success": True,
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "title": core_props.title or "Untitled",
            "author": core_props.author or "Unknown",
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables_data,
            "table_count": len(tables_data),
            "full_text": "\n".join(paragraphs)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


@mcp.tool()
def update_word_document(
    file_path: str,
    action: str,
    content: Optional[str] = None,
    paragraph_index: Optional[int] = None,
    heading_level: Optional[int] = 1
) -> dict:
    """
    Update an existing Word document.
    
    Args:
        file_path: Path to the Word document (.docx file).
                   If relative or just filename, searches in word/ subdirectory first.
        action: Action to perform - "append" (add to end), "insert" (insert at paragraph_index), 
                "replace" (replace paragraph at paragraph_index), "add_heading" (add heading)
        content: Content to add/insert/replace
        paragraph_index: Index of paragraph for insert/replace operations (0-based)
        heading_level: Level of heading (1-9) for add_heading action
    
    Returns:
        Dictionary with operation result
    """
    try:
        # Try to normalize path (searches in word/ subdirectory if relative)
        normalized_path = normalize_file_path(file_path)
        
        # Check both normalized path and original path
        if os.path.exists(normalized_path):
            file_path = normalized_path
        elif not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "file_path": file_path
            }
        
        # Open document
        doc = Document(file_path)
        
        if action == "append":
            # Add content to the end
            if content:
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                    else:
                        doc.add_paragraph()
        
        elif action == "insert":
            # Insert content at specific paragraph index
            if paragraph_index is None:
                return {
                    "success": False,
                    "error": "paragraph_index is required for insert action"
                }
            if content:
                para = doc.paragraphs[min(paragraph_index, len(doc.paragraphs) - 1)]
                para.insert_paragraph_before(content)
        
        elif action == "replace":
            # Replace paragraph at specific index
            if paragraph_index is None:
                return {
                    "success": False,
                    "error": "paragraph_index is required for replace action"
                }
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                para.clear()
                para.add_run(content if content else "")
        
        elif action == "add_heading":
            # Add a heading
            if content:
                doc.add_heading(content, level=min(max(heading_level, 1), 9))
        
        else:
            return {
                "success": False,
                "error": f"Unknown action: {action}. Supported actions: append, insert, replace, add_heading"
            }
        
        # Save document
        doc.save(file_path)
        
        return {
            "success": True,
            "message": f"Document updated successfully",
            "file_path": file_path,
            "action": action,
            "updated_at": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path,
            "action": action
        }


@mcp.tool()
def delete_word_document(file_path: str) -> dict:
    """
    Delete a Word document file.
    
    Args:
        file_path: Path to the Word document (.docx file) to delete.
                   If relative or just filename, searches in word/ subdirectory first.
    
    Returns:
        Dictionary with operation result
    """
    try:
        # Try to normalize path (searches in word/ subdirectory if relative)
        normalized_path = normalize_file_path(file_path)
        
        # Check both normalized path and original path
        if os.path.exists(normalized_path):
            file_path = normalized_path
        elif not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "file_path": file_path
            }
        
        # Delete file
        os.remove(file_path)
        
        return {
            "success": True,
            "message": f"Document deleted successfully",
            "file_path": file_path,
            "deleted_at": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


@mcp.tool()
def list_word_documents(directory: str) -> dict:
    """
    List all Word documents (.docx files) in a directory.
    
    Args:
        directory: Directory path to search for Word documents
    
    Returns:
        Dictionary containing list of Word documents with metadata
    """
    try:
        if not os.path.exists(directory):
            return {
                "success": False,
                "error": f"Directory not found: {directory}",
                "directory": directory
            }
        
        if not os.path.isdir(directory):
            return {
                "success": False,
                "error": f"Path is not a directory: {directory}",
                "directory": directory
            }
        
        # Find all .docx files
        docx_files = []
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.endswith('.docx'):
                    file_path = os.path.join(root, file)
                    file_stat = os.stat(file_path)
                    docx_files.append({
                        "name": file,
                        "path": file_path,
                        "size": file_stat.st_size,
                        "created": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                        "modified": datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                    })
        
        return {
            "success": True,
            "directory": directory,
            "count": len(docx_files),
            "documents": docx_files
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "directory": directory
        }


@mcp.tool()
def add_table_to_document(
    file_path: str,
    table_data: List[List[str]],
    title: Optional[str] = None
) -> dict:
    """
    Add a table to a Word document.
    
    Args:
        file_path: Path to the Word document (.docx file).
                   If relative or just filename, searches in word/ subdirectory first.
        table_data: 2D list representing table data (rows and columns)
        title: Optional title/heading before the table
    
    Returns:
        Dictionary with operation result
    """
    try:
        # Try to normalize path (searches in word/ subdirectory if relative)
        normalized_path = normalize_file_path(file_path)
        
        # Check both normalized path and original path
        if os.path.exists(normalized_path):
            file_path = normalized_path
        elif not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "file_path": file_path
            }
        
        # Open document
        doc = Document(file_path)
        
        # Add title if provided
        if title:
            doc.add_heading(title, level=2)
        
        # Create table
        if table_data:
            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 0
            
            if rows > 0 and cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                
                # Populate table
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = str(cell_data)
        
        # Save document
        doc.save(file_path)
        
        return {
            "success": True,
            "message": f"Table added successfully",
            "file_path": file_path,
            "table_rows": rows if table_data else 0,
            "table_cols": cols if table_data else 0,
            "updated_at": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "file_path": file_path
        }


# ==================== Resources ====================

@mcp.resource("file://word_documents")
def list_documents_resource() -> str:
    """
    Get list of Word documents in the word/ subdirectory.
    """
    try:
        word_dir = get_default_word_dir()
        if os.path.exists(word_dir):
            docx_files = [f for f in os.listdir(word_dir) if f.endswith('.docx')]
        else:
            docx_files = []
        return json.dumps({
            "directory": word_dir,
            "count": len(docx_files),
            "documents": docx_files
        }, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, indent=2)


# ==================== Prompts ====================

@mcp.prompt()
def word_document_help() -> List[Dict]:
    """
    Get help for using Word document operations.
    """
    return [
        {
            "role": "user",
            "content": """I need help with Word document operations. Please explain:
1. How to create a new Word document
2. How to read content from a Word document
3. How to update/modify a Word document
4. How to delete a Word document
5. How to list all Word documents in a directory
6. How to add tables to documents"""
        }
    ]


# The `mcp` variable is automatically detected by MCP CLI when using `mcp run main.py`
# No need to export a separate server object - `mcp` is already the correct type

if __name__ == "__main__":
    import sys
    
    # Check if running in test mode
    if len(sys.argv) > 1 and sys.argv[1] == "--test":
        # Test mode: verify the server can be imported and initialized
        print("✓ Word MCP Server initialized successfully")
        print(f"✓ Default word directory: {get_default_word_dir()}")
        print(f"✓ Server type: {type(mcp).__module__}.{type(mcp).__name__}")
        print("\nNote: This server is designed to be used with MCP clients (e.g., openMCP).")
        print("Do not run it directly. Use 'mcp run main.py' or connect via MCP client.")
        sys.exit(0)
    
    # Run the MCP server using stdio transport (standard for MCP)
    # This expects JSON-RPC messages via stdin, so it will error if run directly
    # Use with MCP clients like openMCP instead
    mcp.run()

