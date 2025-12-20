"""
Core document management functionality.
"""

import os
from pathlib import Path
from typing import List, Dict, Optional, Union, Any
from datetime import datetime
from functools import lru_cache
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .exceptions import (
    DocumentError,
    DocumentNotFoundError,
    DocumentOperationError,
    InvalidPathError,
    ImageError
)
from .path_utils import PathUtils

# Import config - handle both package and direct execution
try:
    from ..config import config
except ImportError:
    from config import config

logger = logging.getLogger(__name__)


class DocumentManager:
    """Manager for Word document operations with enhanced functionality."""
    
    def __init__(self):
        """Initialize DocumentManager."""
        self.path_utils = PathUtils()
    
    def create_document(
        self,
        file_path: Optional[Union[str, Path]] = None,
        title: Optional[str] = None,
        content: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create a new Word document.
        
        Args:
            file_path: File path or filename
            title: Optional title for the document
            content: Optional initial content
        
        Returns:
            Dictionary with operation result
            
        Raises:
            DocumentOperationError: If creation fails
        """
        try:
            # Generate default filename if not provided
            if not file_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                file_path = f"document_{timestamp}.docx"
            
            # Normalize and validate path
            path = self.path_utils.normalize_file_path(file_path)
            path = self.path_utils.validate_file_path(
                path,
                must_exist=False,
                check_size=False
            )
            
            # Ensure parent directory exists
            self.path_utils.ensure_parent_directory(path)
            
            logger.info(f"Creating document: {path}")
            
            # Create new document
            doc = Document()
            
            # Add title if provided
            if title:
                title_para = doc.add_heading(title, level=1)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug(f"Added title: {title}")
            
            # Add content if provided
            if content:
                self._add_content_to_document(doc, content)
                logger.debug(f"Added content: {len(content)} characters")
            
            # Save document
            doc.save(str(path))
            
            result = {
                "success": True,
                "message": "Word document created successfully",
                "file_path": str(path),
                "file_size": path.stat().st_size,
                "created_at": datetime.now().isoformat()
            }
            
            logger.info(f"Document created successfully: {path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to create document: {e}", exc_info=True)
            raise DocumentOperationError("create", str(e), str(file_path))
    
    def read_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Read content from a Word document.
        
        Args:
            file_path: Path to the Word document
        
        Returns:
            Dictionary containing document content and metadata
            
        Raises:
            DocumentNotFoundError: If file not found
            DocumentOperationError: If reading fails
        """
        try:
            # Resolve file path
            path = self.path_utils.resolve_file_path(str(file_path))
            path = self.path_utils.validate_file_path(path, must_exist=True)
            
            logger.info(f"Reading document: {path}")
            
            # Open document
            doc = Document(str(path))
            
            # Extract paragraphs
            paragraphs = [
                para.text.strip()
                for para in doc.paragraphs
                if para.text.strip()
            ]
            
            # Extract tables
            tables_data = self._extract_tables(doc)
            
            # Get document properties
            core_props = doc.core_properties
            
            result = {
                "success": True,
                "file_path": str(path),
                "file_size": path.stat().st_size,
                "title": core_props.title or "Untitled",
                "author": core_props.author or "Unknown",
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "tables": tables_data,
                "table_count": len(tables_data),
                "full_text": "\n".join(paragraphs)
            }
            
            logger.info(
                f"Document read successfully: {path} "
                f"({len(paragraphs)} paragraphs, {len(tables_data)} tables)"
            )
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to read document: {e}", exc_info=True)
            raise DocumentOperationError("read", str(e), str(file_path))
    
    def update_document(
        self,
        file_path: Union[str, Path],
        action: str,
        content: Optional[str] = None,
        paragraph_index: Optional[int] = None,
        heading_level: int = 1
    ) -> Dict[str, Any]:
        """
        Update an existing Word document.
        
        Args:
            file_path: Path to the Word document
            action: Action to perform (append, insert, replace, add_heading)
            content: Content to add/insert/replace
            paragraph_index: Index for insert/replace operations
            heading_level: Level of heading (1-9)
        
        Returns:
            Dictionary with operation result
            
        Raises:
            DocumentNotFoundError: If file not found
            DocumentOperationError: If update fails
        """
        try:
            # Resolve file path
            path = self.path_utils.resolve_file_path(str(file_path))
            path = self.path_utils.validate_file_path(path, must_exist=True)
            
            logger.info(f"Updating document: {path} (action={action})")
            
            # Open document
            doc = Document(str(path))
            
            # Perform action
            if action == "append":
                self._append_content(doc, content)
            elif action == "insert":
                self._insert_content(doc, content, paragraph_index)
            elif action == "replace":
                self._replace_content(doc, content, paragraph_index)
            elif action == "add_heading":
                self._add_heading(doc, content, heading_level)
            else:
                raise InvalidPathError(
                    f"Unknown action: {action}. "
                    f"Supported: append, insert, replace, add_heading"
                )
            
            # Save document
            doc.save(str(path))
            
            result = {
                "success": True,
                "message": "Document updated successfully",
                "file_path": str(path),
                "action": action,
                "updated_at": datetime.now().isoformat()
            }
            
            logger.info(f"Document updated successfully: {path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to update document: {e}", exc_info=True)
            raise DocumentOperationError("update", str(e), str(file_path))
    
    def delete_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Delete a Word document file.
        
        Args:
            file_path: Path to the document to delete
        
        Returns:
            Dictionary with operation result
            
        Raises:
            DocumentNotFoundError: If file not found
            DocumentOperationError: If deletion fails
        """
        try:
            # Resolve file path
            path = self.path_utils.resolve_file_path(str(file_path))
            
            logger.info(f"Deleting document: {path}")
            
            # Delete file
            path.unlink()
            
            result = {
                "success": True,
                "message": "Document deleted successfully",
                "file_path": str(path),
                "deleted_at": datetime.now().isoformat()
            }
            
            logger.info(f"Document deleted successfully: {path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to delete document: {e}", exc_info=True)
            raise DocumentOperationError("delete", str(e), str(file_path))
    
    def list_documents(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
        max_depth: int = 3
    ) -> Dict[str, Any]:
        """
        List all Word documents in a directory.
        
        Args:
            directory: Directory path to search
            recursive: Whether to search recursively
            max_depth: Maximum depth for recursive search
        
        Returns:
            Dictionary containing list of documents with metadata
            
        Raises:
            InvalidPathError: If directory invalid
        """
        try:
            dir_path = Path(directory)
            
            if not dir_path.exists():
                raise InvalidPathError(
                    f"Directory not found: {directory}",
                    str(directory)
                )
            
            if not dir_path.is_dir():
                raise InvalidPathError(
                    f"Path is not a directory: {directory}",
                    str(directory)
                )
            
            logger.info(f"Listing documents in: {dir_path}")
            
            # Find .docx files
            docx_files = []
            
            if recursive:
                for file in dir_path.rglob("*.docx"):
                    # Check depth
                    depth = len(file.relative_to(dir_path).parts)
                    if depth <= max_depth:
                        docx_files.append(self._get_file_info(file))
            else:
                for file in dir_path.glob("*.docx"):
                    docx_files.append(self._get_file_info(file))
            
            result = {
                "success": True,
                "directory": str(dir_path),
                "count": len(docx_files),
                "documents": docx_files
            }
            
            logger.info(f"Found {len(docx_files)} documents in: {dir_path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to list documents: {e}", exc_info=True)
            raise DocumentOperationError("list", str(e), str(directory))
    
    def add_table(
        self,
        file_path: Union[str, Path],
        table_data: List[List[str]],
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Add a table to a Word document.
        
        Args:
            file_path: Path to the document
            table_data: 2D list representing table data
            title: Optional title/heading before the table
        
        Returns:
            Dictionary with operation result
            
        Raises:
            DocumentNotFoundError: If file not found
            DocumentOperationError: If operation fails
        """
        try:
            # Resolve file path
            path = self.path_utils.resolve_file_path(str(file_path))
            path = self.path_utils.validate_file_path(path, must_exist=True)
            
            logger.info(f"Adding table to document: {path}")
            
            # Open document
            doc = Document(str(path))
            
            # Add title if provided
            if title:
                doc.add_heading(title, level=2)
            
            # Create table
            if table_data:
                rows = len(table_data)
                cols = max(len(row) for row in table_data) if table_data else 0
                
                if rows > 0 and cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Populate table
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            if j < cols:
                                table.rows[i].cells[j].text = str(cell_data)
            
            # Save document
            doc.save(str(path))
            
            result = {
                "success": True,
                "message": "Table added successfully",
                "file_path": str(path),
                "table_rows": rows if table_data else 0,
                "table_cols": cols if table_data else 0,
                "updated_at": datetime.now().isoformat()
            }
            
            logger.info(f"Table added successfully: {path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to add table: {e}", exc_info=True)
            raise DocumentOperationError("add_table", str(e), str(file_path))
    
    def insert_image(
        self,
        file_path: Union[str, Path],
        image_path: Union[str, Path],
        width: Optional[float] = None,
        caption: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Insert an image into a Word document.
        
        Args:
            file_path: Path to the document
            image_path: Path to the image file
            width: Width of image in inches (maintains aspect ratio)
            caption: Optional caption text
        
        Returns:
            Dictionary with operation result
            
        Raises:
            DocumentNotFoundError: If file not found
            ImageError: If image operation fails
        """
        try:
            # Resolve paths
            doc_path = self.path_utils.resolve_file_path(str(file_path))
            doc_path = self.path_utils.validate_file_path(doc_path, must_exist=True)
            
            img_path = self.path_utils.validate_image_path(image_path, must_exist=True)
            
            logger.info(f"Inserting image into document: {doc_path}")
            
            # Open document
            doc = Document(str(doc_path))
            
            # Add image
            if width:
                doc.add_picture(str(img_path), width=Inches(width))
            else:
                doc.add_picture(str(img_path))
            
            # Add caption if provided
            if caption:
                para = doc.add_paragraph(caption)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].font.italic = True
                para.runs[0].font.size = Pt(10)
            
            # Save document
            doc.save(str(doc_path))
            
            result = {
                "success": True,
                "message": "Image inserted successfully",
                "file_path": str(doc_path),
                "image_path": str(img_path),
                "updated_at": datetime.now().isoformat()
            }
            
            logger.info(f"Image inserted successfully: {doc_path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Failed to insert image: {e}", exc_info=True)
            raise ImageError(str(e), str(image_path))
    
    # Helper methods
    
    @staticmethod
    def _add_content_to_document(doc: Document, content: str) -> None:
        """Add content to document, handling newlines."""
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph()
    
    @staticmethod
    def _extract_tables(doc: Document) -> List[List[List[str]]]:
        """Extract tables from document."""
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        return tables_data
    
    @staticmethod
    def _append_content(doc: Document, content: Optional[str]) -> None:
        """Append content to document."""
        if content:
            DocumentManager._add_content_to_document(doc, content)
    
    @staticmethod
    def _insert_content(
        doc: Document,
        content: Optional[str],
        paragraph_index: Optional[int]
    ) -> None:
        """Insert content at specific paragraph index."""
        if paragraph_index is None:
            raise DocumentOperationError(
                "insert",
                "paragraph_index is required for insert action"
            )
        if content:
            para = doc.paragraphs[min(paragraph_index, len(doc.paragraphs) - 1)]
            para.insert_paragraph_before(content)
    
    @staticmethod
    def _replace_content(
        doc: Document,
        content: Optional[str],
        paragraph_index: Optional[int]
    ) -> None:
        """Replace paragraph at specific index."""
        if paragraph_index is None:
            raise DocumentOperationError(
                "replace",
                "paragraph_index is required for replace action"
            )
        if paragraph_index < len(doc.paragraphs):
            para = doc.paragraphs[paragraph_index]
            para.clear()
            para.add_run(content if content else "")
    
    @staticmethod
    def _add_heading(
        doc: Document,
        content: Optional[str],
        heading_level: int
    ) -> None:
        """Add a heading to document."""
        if content:
            level = min(max(heading_level, 1), 9)
            doc.add_heading(content, level=level)
    
    @staticmethod
    def _get_file_info(file_path: Path) -> Dict[str, Any]:
        """Get file information."""
        stat = file_path.stat()
        return {
            "name": file_path.name,
            "path": str(file_path),
            "size": stat.st_size,
            "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
        }

