"""
Word Document MCP Server

A simple MCP server for Word document operations including:
- Create, read, update, delete documents
- Add tables and images
- Format text and search/replace
- Merge documents
"""

from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import httpx
import re
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Any

# Initialize FastMCP server
mcp = FastMCP("Word Document MCP Server")

# Default word directory
WORD_DIR = Path("word")
WORD_DIR.mkdir(exist_ok=True)

# Load config
def load_config() -> dict:
    """Load configuration from mcpconfig.json"""
    config_path = Path(__file__).parent / "mcpconfig.json"
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

CONFIG = load_config()
GOOGLE_API_KEY = CONFIG.get("google", "")


# ==================== Helper Functions ====================

# 工具信息映射，用于任务规划
TOOL_REGISTRY: Dict[str, Dict[str, Any]] = {
    "create_document": {
        "description": "创建新文档",
        "keywords": ["创建", "新建", "生成", "写", "create", "new"],
        "params": ["filename", "title", "content"],
        "required": []
    },
    "read_document": {
        "description": "读取文档内容",
        "keywords": ["读取", "查看", "打开", "获取", "read", "open", "get"],
        "params": ["filename"],
        "required": ["filename"]
    },
    "update_document": {
        "description": "更新文档（追加、插入、替换内容）",
        "keywords": ["更新", "修改", "追加", "添加内容", "update", "append", "modify"],
        "params": ["filename", "action", "content", "paragraph_index"],
        "required": ["filename", "action"]
    },
    "delete_document": {
        "description": "删除文档",
        "keywords": ["删除", "移除", "delete", "remove"],
        "params": ["filename"],
        "required": ["filename"]
    },
    "list_documents": {
        "description": "列出所有文档",
        "keywords": ["列出", "显示", "所有文档", "list", "show all"],
        "params": [],
        "required": []
    },
    "add_table": {
        "description": "添加表格",
        "keywords": ["表格", "table", "添加表", "插入表"],
        "params": ["filename", "table_data", "title"],
        "required": ["filename", "table_data"]
    },
    "insert_image": {
        "description": "插入图片",
        "keywords": ["图片", "插入图", "添加图", "image", "picture", "photo"],
        "params": ["filename", "image_path", "width"],
        "required": ["filename", "image_path"]
    },
    "format_text": {
        "description": "格式化文本（加粗、斜体、字号）",
        "keywords": ["格式", "加粗", "斜体", "字体", "format", "bold", "italic"],
        "params": ["filename", "paragraph_index", "bold", "italic", "font_size"],
        "required": ["filename", "paragraph_index"]
    },
    "search_replace": {
        "description": "搜索替换文本",
        "keywords": ["替换", "搜索", "查找", "replace", "search", "find"],
        "params": ["filename", "search_text", "replace_text"],
        "required": ["filename", "search_text", "replace_text"]
    },
    "download_image": {
        "description": "从URL下载图片",
        "keywords": ["下载图片", "下载图", "download image"],
        "params": ["url", "filename"],
        "required": ["url"]
    },
    "google_search": {
        "description": "Google搜索获取信息",
        "keywords": ["搜索", "查询", "search", "google", "查资料"],
        "params": ["query", "num_results"],
        "required": ["query"]
    },
    "google_image_search": {
        "description": "Google图片搜索",
        "keywords": ["搜图", "搜索图片", "找图", "image search"],
        "params": ["query", "num_results"],
        "required": ["query"]
    }
}


def extract_filename(text: str) -> Optional[str]:
    """从文本中提取文件名"""
    # 匹配 .docx 文件名
    match = re.search(r'[\w\u4e00-\u9fff_-]+\.docx', text, re.IGNORECASE)
    if match:
        return match.group()
    
    # 匹配引号中的文件名
    match = re.search(r'["\']([^"\']+)["\']', text)
    if match:
        return match.group(1)
    
    # 匹配"文档X"、"文件X"模式
    match = re.search(r'(?:文档|文件|document)\s*[：:]*\s*([\w\u4e00-\u9fff_-]+)', text, re.IGNORECASE)
    if match:
        return match.group(1)
    
    return None


def extract_title(text: str) -> Optional[str]:
    """从文本中提取标题"""
    patterns = [
        r'标题[：:为是]\s*["\']?([^"\'，。\n]+)["\']?',
        r'题目[：:为是]\s*["\']?([^"\'，。\n]+)["\']?',
        r'title[：:]\s*["\']?([^"\'，。\n]+)["\']?',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return None


def extract_content(text: str) -> Optional[str]:
    """从文本中提取内容"""
    patterns = [
        r'内容[：:为是]\s*["\']?(.+?)["\']?(?:[，。]|$)',
        r'content[：:]\s*["\']?(.+?)["\']?(?:[，。]|$)',
        r'写[：:]\s*["\']?(.+?)["\']?(?:[，。]|$)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            return match.group(1).strip()
    return None


def match_intent(text: str) -> List[Dict[str, Any]]:
    """匹配用户意图，返回可能的工具列表（按匹配度排序）"""
    text_lower = text.lower()
    scores = []
    
    for tool_name, info in TOOL_REGISTRY.items():
        score = 0
        matched_keywords = []
        
        for keyword in info["keywords"]:
            if keyword in text_lower:
                score += len(keyword)  # 关键词越长，匹配越精确
                matched_keywords.append(keyword)
        
        if score > 0:
            scores.append({
                "tool": tool_name,
                "score": score,
                "matched_keywords": matched_keywords,
                "description": info["description"],
                "required_params": info["required"],
                "all_params": info["params"]
            })
    
    # 按分数降序排序
    scores.sort(key=lambda x: x["score"], reverse=True)
    return scores


def get_file_path(filename: str) -> Path:
    """Get full path for a file in word directory."""
    path = Path(filename)
    if not path.is_absolute():
        path = WORD_DIR / filename
    if not str(path).endswith('.docx'):
        path = Path(str(path) + '.docx')
    return path


# ==================== Tools ====================

@mcp.tool()
def plan_task(user_input: str) -> dict:
    """
    智能任务规划工具：解析用户的非结构化输入，识别意图并生成结构化执行计划。
    
    当用户输入模糊或复杂的自然语言描述时，先调用此工具进行解析，
    然后根据返回的执行计划依次调用对应的工具。
    
    Args:
        user_input: 用户的自然语言输入，如 "创建一个标题为年度报告的文档，包含销售数据表格"
    
    Returns:
        结构化的执行计划，包含:
        - success: 是否成功解析
        - intent_summary: 意图摘要
        - steps: 执行步骤列表，每步包含 tool_name, params, description
        - extracted_info: 从输入中提取的信息
        - confidence: 置信度 (high/medium/low)
        - suggestions: 如果信息不完整，给出建议
    """
    try:
        # 1. 匹配意图
        matched_intents = match_intent(user_input)
        
        if not matched_intents:
            return {
                "success": False,
                "error": "无法识别用户意图",
                "suggestions": [
                    "请明确说明您想要的操作，例如：",
                    "- 创建文档：'创建一个名为xxx的文档'",
                    "- 读取文档：'打开/读取xxx文档'",
                    "- 添加表格：'在文档中添加表格'",
                    "- 插入图片：'在文档中插入图片'",
                    "- 搜索信息：'搜索关于xxx的信息'"
                ],
                "available_tools": list(TOOL_REGISTRY.keys())
            }
        
        # 2. 提取参数
        extracted = {
            "filename": extract_filename(user_input),
            "title": extract_title(user_input),
            "content": extract_content(user_input)
        }
        
        # 3. 生成执行计划
        steps = []
        primary_intent = matched_intents[0]
        
        # 检查是否是复合任务（多个意图）
        is_complex = len(matched_intents) > 1 and matched_intents[1]["score"] > 2
        
        # 生成主要步骤
        step1_params = {}
        missing_params = []
        
        for param in primary_intent["required_params"]:
            if param == "filename" and extracted["filename"]:
                step1_params["filename"] = extracted["filename"]
            elif param == "title" and extracted["title"]:
                step1_params["title"] = extracted["title"]
            elif param == "content" and extracted["content"]:
                step1_params["content"] = extracted["content"]
            elif param == "query":
                # 对于搜索，使用整个输入作为查询（去掉动词）
                query = re.sub(r'^(搜索|查询|查找|找|search)\s*', '', user_input, flags=re.IGNORECASE)
                step1_params["query"] = query.strip() or user_input
            else:
                missing_params.append(param)
        
        # 添加可选参数
        for param in primary_intent["all_params"]:
            if param not in step1_params:
                if param == "title" and extracted["title"]:
                    step1_params["title"] = extracted["title"]
                elif param == "content" and extracted["content"]:
                    step1_params["content"] = extracted["content"]
        
        steps.append({
            "step": 1,
            "tool_name": primary_intent["tool"],
            "description": primary_intent["description"],
            "params": step1_params,
            "missing_params": missing_params
        })
        
        # 如果是复合任务，添加后续步骤
        if is_complex:
            for i, intent in enumerate(matched_intents[1:3], start=2):  # 最多3个步骤
                if intent["score"] > 2:
                    step_params = {}
                    if "filename" in intent["all_params"] and extracted["filename"]:
                        step_params["filename"] = extracted["filename"]
                    elif "filename" in intent["all_params"] and steps[0]["params"].get("filename"):
                        step_params["filename"] = steps[0]["params"]["filename"]
                    
                    steps.append({
                        "step": i,
                        "tool_name": intent["tool"],
                        "description": intent["description"],
                        "params": step_params,
                        "missing_params": [p for p in intent["required_params"] if p not in step_params]
                    })
        
        # 4. 计算置信度
        has_all_required = len(missing_params) == 0
        confidence = "high" if has_all_required and primary_intent["score"] > 3 else \
                    "medium" if primary_intent["score"] > 2 else "low"
        
        # 5. 生成建议
        suggestions = []
        if missing_params:
            suggestions.append(f"缺少必要参数: {', '.join(missing_params)}")
            if "filename" in missing_params:
                suggestions.append("请提供文档名称，例如：'报告.docx' 或 '我的文档'")
            if "table_data" in missing_params:
                suggestions.append("请提供表格数据，格式为二维数组")
            if "image_path" in missing_params:
                suggestions.append("请提供图片路径，或先使用 download_image 下载图片")
        
        return {
            "success": True,
            "intent_summary": f"识别到主要意图: {primary_intent['description']}",
            "matched_keywords": primary_intent["matched_keywords"],
            "confidence": confidence,
            "is_complex_task": is_complex,
            "steps": steps,
            "extracted_info": {k: v for k, v in extracted.items() if v},
            "suggestions": suggestions if suggestions else ["参数完整，可以直接执行"],
            "raw_matches": matched_intents[:3]  # 返回前3个匹配结果供参考
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"解析出错: {str(e)}"
        }


@mcp.tool()
def get_tool_info(tool_name: Optional[str] = None) -> dict:
    """
    获取工具信息。不带参数返回所有工具列表，带参数返回特定工具的详细信息。
    
    Args:
        tool_name: 工具名称（可选）
    
    Returns:
        工具信息
    """
    if tool_name:
        if tool_name in TOOL_REGISTRY:
            info = TOOL_REGISTRY[tool_name]
            return {
                "success": True,
                "tool_name": tool_name,
                "description": info["description"],
                "parameters": info["params"],
                "required_parameters": info["required"],
                "trigger_keywords": info["keywords"]
            }
        else:
            return {
                "success": False,
                "error": f"未知工具: {tool_name}",
                "available_tools": list(TOOL_REGISTRY.keys())
            }
    else:
        tools_summary = []
        for name, info in TOOL_REGISTRY.items():
            tools_summary.append({
                "name": name,
                "description": info["description"],
                "required_params": info["required"]
            })
        return {
            "success": True,
            "tools_count": len(tools_summary),
            "tools": tools_summary
        }


@mcp.tool()
def create_document(
    filename: Optional[str] = None,
    title: Optional[str] = None,
    content: Optional[str] = None
) -> dict:
    """
    Create a new Word document.
    
    Args:
        filename: File name (optional, auto-generated if not provided)
        title: Document title (optional)
        content: Initial content (optional)
    
    Returns:
        Result dictionary with file path and status
    """
    try:
        # Generate filename if not provided
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        file_path = get_file_path(filename)
        
        # Create document
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if content:
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Document created successfully",
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def read_document(filename: str) -> dict:
    """
    Read content from a Word document.
    
    Args:
        filename: File name to read
    
    Returns:
        Document content and metadata
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        # Extract paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(table_data)
        
        return {
            "success": True,
            "file_path": str(file_path),
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables),
            "full_text": "\n".join(paragraphs)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def update_document(
    filename: str,
    action: str,
    content: Optional[str] = None,
    paragraph_index: Optional[int] = None
) -> dict:
    """
    Update a Word document.
    
    Args:
        filename: File name to update
        action: Action type - "append", "insert", "replace", "add_heading"
        content: Content to add/insert/replace
        paragraph_index: Paragraph index for insert/replace (0-based)
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        if action == "append":
            for line in content.split('\n') if content else []:
                if line.strip():
                    doc.add_paragraph(line.strip())
        
        elif action == "add_heading":
            if content:
                doc.add_heading(content, level=2)
        
        elif action == "insert" and paragraph_index is not None:
            if content and paragraph_index < len(doc.paragraphs):
                doc.paragraphs[paragraph_index].insert_paragraph_before(content)
        
        elif action == "replace" and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                para.clear()
                para.add_run(content or "")
        
        else:
            return {"success": False, "error": f"Invalid action or missing parameters"}
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Document updated successfully",
            "action": action
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def delete_document(filename: str) -> dict:
    """
    Delete a Word document.
    
    Args:
        filename: File name to delete
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        file_path.unlink()
        
        return {
            "success": True,
            "message": "Document deleted successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def list_documents() -> dict:
    """
    List all Word documents in the word directory.
    
    Returns:
        List of documents with metadata
    """
    try:
        docs = []
        for file in WORD_DIR.glob("*.docx"):
            stat = file.stat()
            docs.append({
                "name": file.name,
                "path": str(file),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
        
        return {
            "success": True,
            "count": len(docs),
            "documents": docs
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def add_table(
    filename: str,
    table_data: List[List[str]],
    title: Optional[str] = None
) -> dict:
    """
    Add a table to a document.
    
    Args:
        filename: File name
        table_data: 2D list of table data
        title: Optional table title
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        if title:
            doc.add_heading(title, level=2)
        
        if table_data:
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < cols:
                        table.rows[i].cells[j].text = str(cell_data)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Table added successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def insert_image(
    filename: str,
    image_path: str,
    width: Optional[float] = None
) -> dict:
    """
    Insert an image into a document.
    
    Args:
        filename: Document file name
        image_path: Path to image file
        width: Image width in inches (optional)
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        if not Path(image_path).exists():
            return {"success": False, "error": f"Image not found: {image_path}"}
        
        doc = Document(str(file_path))
        
        if width:
            doc.add_picture(image_path, width=Inches(width))
        else:
            doc.add_picture(image_path)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Image inserted successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def format_text(
    filename: str,
    paragraph_index: int,
    bold: bool = False,
    italic: bool = False,
    font_size: Optional[int] = None
) -> dict:
    """
    Format text in a paragraph.
    
    Args:
        filename: Document file name
        paragraph_index: Paragraph index (0-based)
        bold: Make text bold
        italic: Make text italic
        font_size: Font size in points
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        if paragraph_index >= len(doc.paragraphs):
            return {"success": False, "error": "Paragraph index out of range"}
        
        para = doc.paragraphs[paragraph_index]
        
        for run in para.runs:
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if font_size:
                run.font.size = Pt(font_size)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Text formatted successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def search_replace(
    filename: str,
    search_text: str,
    replace_text: str
) -> dict:
    """
    Search and replace text in a document.
    
    Args:
        filename: Document file name
        search_text: Text to search for
        replace_text: Text to replace with
    
    Returns:
        Result with replacement count
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        count = 0
        
        for para in doc.paragraphs:
            for run in para.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)
                    count += run.text.count(replace_text)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": f"Replaced {count} occurrences",
            "replacement_count": count
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def download_image(
    url: str,
    filename: Optional[str] = None
) -> dict:
    """
    从 URL 下载图片到本地。下载后可使用 insert_image 工具将图片插入文档。
    
    Args:
        url: 图片的 URL 地址
        filename: 保存的文件名（可选，不含扩展名，自动根据 URL 生成）
    
    Returns:
        包含本地图片路径的结果
    """
    try:
        # 创建图片目录
        images_dir = WORD_DIR / "images"
        images_dir.mkdir(exist_ok=True)
        
        # 生成文件名
        if not filename:
            # 从 URL 提取文件名或生成时间戳文件名
            url_filename = url.split("/")[-1].split("?")[0]
            if url_filename and "." in url_filename:
                filename = url_filename
            else:
                filename = f"image_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
        else:
            # 确保有扩展名
            if not any(filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']):
                filename = f"{filename}.jpg"
        
        file_path = images_dir / filename
        
        # 下载图片
        with httpx.Client(timeout=30.0, follow_redirects=True) as client:
            response = client.get(url)
            response.raise_for_status()
            
            # 检查是否是图片
            content_type = response.headers.get("content-type", "")
            if not content_type.startswith("image/"):
                return {"success": False, "error": f"URL 不是图片: {content_type}"}
            
            # 保存图片
            with open(file_path, "wb") as f:
                f.write(response.content)
        
        return {
            "success": True,
            "message": "图片下载成功",
            "local_path": str(file_path),
            "filename": filename,
            "size": file_path.stat().st_size
        }
        
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"下载失败: HTTP {e.response.status_code}"}
    except Exception as e:
        return {"success": False, "error": f"下载出错: {str(e)}"}


@mcp.tool()
def google_image_search(
    query: str,
    num_results: int = 5
) -> dict:
    """
    使用 Google 搜索图片。返回图片的 URL 列表，可配合 download_image 下载后插入文档。
    
    Args:
        query: 搜索关键词
        num_results: 返回结果数量，默认5条
    
    Returns:
        图片搜索结果，包含图片 URL、标题和来源
    """
    if not GOOGLE_API_KEY:
        return {"success": False, "error": "Google API Key 未配置"}
    
    try:
        # 使用 Serper.dev 图片搜索 API
        url = "https://google.serper.dev/images"
        headers = {
            "X-API-KEY": GOOGLE_API_KEY,
            "Content-Type": "application/json"
        }
        payload = {
            "q": query,
            "num": num_results,
            "hl": "zh-CN"
        }
        
        with httpx.Client(timeout=30.0) as client:
            response = client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
        
        results = []
        images_results = data.get("images", [])
        
        for item in images_results[:num_results]:
            results.append({
                "title": item.get("title", ""),
                "image_url": item.get("imageUrl", ""),
                "thumbnail": item.get("thumbnailUrl", ""),
                "source": item.get("source", ""),
                "link": item.get("link", "")
            })
        
        return {
            "success": True,
            "query": query,
            "count": len(results),
            "images": results
        }
        
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"搜索请求失败: {e.response.status_code}"}
    except Exception as e:
        return {"success": False, "error": f"搜索出错: {str(e)}"}


@mcp.tool()
def google_search(
    query: str,
    num_results: int = 5
) -> dict:
    """
    使用 Google 搜索获取信息。当需要查询最新资讯、获取某个主题的详细信息时使用此工具。
    
    Args:
        query: 搜索关键词
        num_results: 返回结果数量，默认5条
    
    Returns:
        搜索结果列表，包含标题、链接和摘要
    """
    if not GOOGLE_API_KEY:
        return {"success": False, "error": "Google API Key 未配置"}
    
    try:
        # 使用 Serper.dev API
        url = "https://google.serper.dev/search"
        headers = {
            "X-API-KEY": GOOGLE_API_KEY,
            "Content-Type": "application/json"
        }
        payload = {
            "q": query,
            "num": num_results,
            "hl": "zh-CN",
            "gl": "cn"
        }
        
        with httpx.Client(timeout=30.0) as client:
            response = client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
        
        results = []
        
        # 提取有机搜索结果
        organic_results = data.get("organic", [])
        for item in organic_results[:num_results]:
            results.append({
                "title": item.get("title", ""),
                "link": item.get("link", ""),
                "snippet": item.get("snippet", "")
            })
        
        # 如果有答案框
        if "answerBox" in data:
            answer = data["answerBox"]
            results.insert(0, {
                "title": answer.get("title", "答案"),
                "link": answer.get("link", ""),
                "snippet": answer.get("answer", answer.get("snippet", ""))
            })
        
        return {
            "success": True,
            "query": query,
            "count": len(results),
            "results": results
        }
        
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"搜索请求失败: {e.response.status_code}"}
    except Exception as e:
        return {"success": False, "error": f"搜索出错: {str(e)}"}


# ==================== Resources ====================

@mcp.resource("file://documents")
def list_documents_resource() -> str:
    """Get list of all Word documents."""
    result = list_documents()
    return json.dumps(result, indent=2)


# ==================== Prompts ====================

@mcp.prompt()
def help_prompt() -> list[dict]:
    """Get help for Word document operations."""
    return [{
        "role": "user",
        "content": """Help me with Word document operations:
1. How to create a document
2. How to add tables
3. How to format text
4. How to search and replace"""
    }]


# Run server
if __name__ == "__main__":
    mcp.run()

