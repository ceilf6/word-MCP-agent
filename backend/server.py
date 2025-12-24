"""
Word MCP Server with SSE Transport + LLM Agent

启动方式:
    python server.py

服务器会在 http://localhost:8080 启动 SSE 端点
使用 mcpconfig.json 中的 defaultLLM 配置调用大模型
"""

import asyncio
import json
import logging
from typing import Optional, List, Any
from pathlib import Path
from datetime import datetime
import re
import httpx

from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 导入多 Agent 模块
from agents import DocumentCreationPipeline, StructurizerAgent

# 导入记忆系统
from memory import MemoryManager, get_session, remember, recall

# 默认字体设置
DEFAULT_FONT_NAME = "宋体"
DEFAULT_FONT_SIZE = Pt(12)


def set_run_font(run, font_name: str = DEFAULT_FONT_NAME, font_size=DEFAULT_FONT_SIZE):
    """设置 run 的字体为宋体"""
    run.font.name = font_name
    run.font.size = font_size
    # 设置中文字体（东亚字体）
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI 应用
app = FastAPI(title="Word MCP Server")

# CORS 配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 文档目录
WORD_DIR = Path("word")
WORD_DIR.mkdir(exist_ok=True)

# ==================== 加载配置 ====================

def load_config() -> dict:
    """从 mcpconfig.json 加载所有配置"""
    config_path = Path(__file__).parent / "mcpconfig.json"
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"加载配置失败: {e}")
        return {}

CONFIG = load_config()
LLM_CONFIG = CONFIG.get("defaultLLM", {})
GOOGLE_API_KEY = CONFIG.get("google", "")

logger.info(f"LLM 配置: baseURL={LLM_CONFIG.get('baseURL')}, model={LLM_CONFIG.get('model')}")
logger.info(f"Google API Key: {'已配置' if GOOGLE_API_KEY else '未配置'}")


# ==================== 工具函数 ====================

def get_file_path(filename: str) -> Path:
    """获取文件完整路径"""
    path = Path(filename)
    if not path.is_absolute():
        path = WORD_DIR / filename
    if not str(path).endswith('.docx'):
        path = Path(str(path) + '.docx')
    return path


# ==================== MCP 工具定义 ====================

TOOLS = {
    "create_document": {
        "description": "创建新的 Word 文档",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "文件名（不含扩展名）"},
                "title": {"type": "string", "description": "文档标题"},
                "content": {"type": "string", "description": "初始内容（可包含多行）"}
            },
            "required": []
        }
    },
    "read_document": {
        "description": "读取 Word 文档内容",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "文件名"}
            },
            "required": ["filename"]
        }
    },
    "update_document": {
        "description": "更新 Word 文档，支持追加内容、添加标题、插入段落、替换段落",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "文件名"},
                "action": {"type": "string", "enum": ["append", "insert", "replace", "add_heading"], "description": "操作类型"},
                "content": {"type": "string", "description": "要写入的内容"},
                "paragraph_index": {"type": "integer", "description": "段落索引（insert/replace 时使用）"}
            },
            "required": ["filename", "action"]
        }
    },
    "delete_document": {
        "description": "删除 Word 文档",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "文件名"}
            },
            "required": ["filename"]
        }
    },
    "list_documents": {
        "description": "列出所有 Word 文档",
        "parameters": {
            "type": "object",
            "properties": {},
            "required": []
        }
    },
    "add_table": {
        "description": "向文档添加表格",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "文件名"},
                "table_data": {"type": "array", "description": "表格数据（二维数组），第一行为表头", "items": {"type": "array", "items": {"type": "string"}}},
                "title": {"type": "string", "description": "表格标题（可选）"}
            },
            "required": ["filename", "table_data"]
        }
    },
    "search_replace": {
        "description": "搜索并替换文档中的文本",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "文件名"},
                "search_text": {"type": "string", "description": "要搜索的文本"},
                "replace_text": {"type": "string", "description": "替换为的文本"}
            },
            "required": ["filename", "search_text", "replace_text"]
        }
    },
    "google_search": {
        "description": "使用 Google 搜索获取信息。当需要查询最新资讯、获取某个主题的详细信息时使用此工具。",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "搜索关键词"},
                "num_results": {"type": "integer", "description": "返回结果数量，默认5条", "default": 5}
            },
            "required": ["query"]
        }
    },
    "google_image_search": {
        "description": "使用 Google 搜索图片。返回图片 URL 列表，可配合 download_image 下载后用 insert_image 插入文档。",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "图片搜索关键词"},
                "num_results": {"type": "integer", "description": "返回结果数量，默认5条", "default": 5}
            },
            "required": ["query"]
        }
    },
    "download_image": {
        "description": "从 URL 下载图片到本地。下载后可使用 insert_image 工具将图片插入文档。",
        "parameters": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "图片的 URL 地址"},
                "filename": {"type": "string", "description": "保存的文件名（可选）"}
            },
            "required": ["url"]
        }
    },
    "insert_image": {
        "description": "将本地图片插入到 Word 文档中。需要先用 download_image 下载图片。",
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Word 文档文件名"},
                "image_path": {"type": "string", "description": "图片的本地路径"},
                "width": {"type": "number", "description": "图片宽度（英寸），可选"}
            },
            "required": ["filename", "image_path"]
        }
    },
    # ==================== 多 Agent 工具 ====================
    "create_document_with_agents": {
        "description": "【推荐】使用多 Agent 流水线创建文档。经过结构化Agent解析 → 创作Agent生成 → 评审Agent评分，不达标自动重写。创建文档时必须使用此工具！",
        "parameters": {
            "type": "object",
            "properties": {
                "user_request": {"type": "string", "description": "用户的原始自然语言请求，如'帮我写一份产品介绍文档'"},
                "auto_confirm": {"type": "boolean", "description": "是否自动确认（True=跳过澄清问题，False=需要澄清时返回问题）", "default": False}
            },
            "required": ["user_request"]
        }
    },
    "structurize_input": {
        "description": "仅使用结构化 Agent 解析用户输入，不创建文档。用于理解用户意图和提取参数。",
        "parameters": {
            "type": "object",
            "properties": {
                "user_input": {"type": "string", "description": "用户的自然语言输入"}
            },
            "required": ["user_input"]
        }
    },
    # ==================== 记忆工具 ====================
    "save_to_memory": {
        "description": "保存重要信息到长期记忆。用于记住用户偏好、重要事实等。",
        "parameters": {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "要记忆的内容"},
                "category": {"type": "string", "enum": ["fact", "preference", "context"], "description": "分类：fact=事实，preference=偏好，context=上下文"},
                "importance": {"type": "number", "description": "重要性 0-1，默认0.5"},
                "tags": {"type": "array", "items": {"type": "string"}, "description": "标签列表"}
            },
            "required": ["content"]
        }
    },
    "recall_memory": {
        "description": "从长期记忆中检索相关信息。",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "检索关键词"},
                "category": {"type": "string", "enum": ["fact", "preference", "context"], "description": "限定分类"},
                "limit": {"type": "integer", "description": "返回数量，默认5"}
            },
            "required": ["query"]
        }
    },
    "get_memory_stats": {
        "description": "获取当前会话的记忆统计信息。",
        "parameters": {
            "type": "object",
            "properties": {},
            "required": []
        }
    }
}


# ==================== 工具实现 ====================

def create_document(filename: str = None, title: str = None, content: str = None) -> dict:
    """创建新文档（默认使用宋体）"""
    try:
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        file_path = get_file_path(filename)
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置标题字体为宋体
            for run in heading.runs:
                set_run_font(run, font_size=Pt(18))
        
        if content:
            for line in content.split('\n'):
                para = doc.add_paragraph()
                run = para.add_run(line.strip() if line.strip() else "")
                set_run_font(run)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "文档创建成功",
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def read_document(filename: str) -> dict:
    """读取文档"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs]
        
        tables = []
        for table in doc.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(table_data)
        
        return {
            "success": True,
            "file_path": str(file_path),
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables),
            "full_text": "\n".join(paragraphs)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def update_document(filename: str, action: str, content: str = None, paragraph_index: int = None) -> dict:
    """更新文档（默认使用宋体）"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        
        if action == "append" and content:
            for line in content.split('\n'):
                para = doc.add_paragraph()
                run = para.add_run(line)
                set_run_font(run)
        elif action == "add_heading" and content:
            heading = doc.add_heading(content, level=2)
            for run in heading.runs:
                set_run_font(run, font_size=Pt(16))
        elif action == "insert" and content and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                new_para = doc.paragraphs[paragraph_index].insert_paragraph_before()
                run = new_para.add_run(content)
                set_run_font(run)
        elif action == "replace" and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                para.clear()
                run = para.add_run(content or "")
                set_run_font(run)
        else:
            return {"success": False, "error": "无效的操作或缺少参数"}
        
        doc.save(str(file_path))
        return {"success": True, "message": "文档更新成功", "action": action}
    except Exception as e:
        return {"success": False, "error": str(e)}


def delete_document(filename: str) -> dict:
    """删除文档"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        file_path.unlink()
        return {"success": True, "message": "文档删除成功"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def list_documents() -> dict:
    """列出所有文档"""
    try:
        docs = []
        for file in WORD_DIR.glob("*.docx"):
            stat = file.stat()
            docs.append({
                "name": file.name,
                "path": str(file),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
        return {"success": True, "count": len(docs), "documents": docs}
    except Exception as e:
        return {"success": False, "error": str(e)}


def add_table(filename: str, table_data: list, title: str = None) -> dict:
    """添加表格（默认使用宋体）"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        
        if title:
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                set_run_font(run, font_size=Pt(16))
        
        if table_data:
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < cols:
                        cell = table.rows[i].cells[j]
                        cell.text = ""  # 清空默认文本
                        para = cell.paragraphs[0]
                        run = para.add_run(str(cell_data))
                        set_run_font(run)
        
        doc.save(str(file_path))
        return {"success": True, "message": "表格添加成功"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def search_replace(filename: str, search_text: str, replace_text: str) -> dict:
    """搜索替换"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        count = 0
        
        for para in doc.paragraphs:
            for run in para.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)
                    count += 1
        
        doc.save(str(file_path))
        return {"success": True, "message": f"替换了 {count} 处", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


async def google_search_async(query: str, num_results: int = 5) -> dict:
    """
    使用 Serper.dev API 进行 Google 搜索（异步版本）
    需要 GOOGLE_API_KEY 配置（从 serper.dev 获取）
    """
    if not GOOGLE_API_KEY:
        return {"success": False, "error": "Google API Key 未配置"}
    
    try:
        # 使用 Serper.dev API
        url = "https://google.serper.dev/search"
        headers = {
            "X-API-KEY": GOOGLE_API_KEY,
            "Content-Type": "application/json"
        }
        payload = {
            "q": query,
            "num": num_results,
            "hl": "zh-CN",
            "gl": "cn"
        }
        
        async with httpx.AsyncClient(timeout=30.0, proxy=None, trust_env=False) as client:
            response = await client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
        
        # 解析搜索结果
        results = []
        
        # 提取有机搜索结果
        organic_results = data.get("organic", [])
        for item in organic_results[:num_results]:
            results.append({
                "title": item.get("title", ""),
                "link": item.get("link", ""),
                "snippet": item.get("snippet", "")
            })
        
        # 如果有答案框
        if "answerBox" in data:
            answer = data["answerBox"]
            results.insert(0, {
                "title": answer.get("title", "答案"),
                "link": answer.get("link", ""),
                "snippet": answer.get("answer", answer.get("snippet", ""))
            })
        
        return {
            "success": True,
            "query": query,
            "count": len(results),
            "results": results
        }
        
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"搜索请求失败: {e.response.status_code}"}
    except Exception as e:
        return {"success": False, "error": f"搜索出错: {str(e)}"}


def google_search(query: str, num_results: int = 5) -> dict:
    """
    使用 Google 搜索（同步包装器）
    """
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, google_search_async(query, num_results))
                return future.result()
        else:
            return asyncio.run(google_search_async(query, num_results))
    except Exception as e:
        return {"success": False, "error": f"搜索执行失败: {str(e)}"}


def google_image_search(query: str, num_results: int = 5) -> dict:
    """
    使用 Serper.dev API 搜索图片
    """
    if not GOOGLE_API_KEY:
        return {"success": False, "error": "Google API Key 未配置"}
    
    try:
        # 使用 Serper.dev 图片搜索 API
        url = "https://google.serper.dev/images"
        headers = {
            "X-API-KEY": GOOGLE_API_KEY,
            "Content-Type": "application/json"
        }
        payload = {
            "q": query,
            "num": num_results,
            "hl": "zh-CN"
        }
        
        with httpx.Client(timeout=30.0, proxy=None, trust_env=False) as client:
            response = client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
        
        results = []
        images_results = data.get("images", [])
        
        for item in images_results[:num_results]:
            results.append({
                "title": item.get("title", ""),
                "image_url": item.get("imageUrl", ""),
                "thumbnail": item.get("thumbnailUrl", ""),
                "source": item.get("source", ""),
                "link": item.get("link", "")
            })
        
        return {
            "success": True,
            "query": query,
            "count": len(results),
            "images": results
        }
        
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"搜索请求失败: {e.response.status_code}"}
    except Exception as e:
        return {"success": False, "error": f"搜索出错: {str(e)}"}


def download_image(url: str, filename: str = None) -> dict:
    """
    从 URL 下载图片到本地
    """
    try:
        # 创建图片目录
        images_dir = WORD_DIR / "images"
        images_dir.mkdir(exist_ok=True)
        
        # 生成文件名
        if not filename:
            url_filename = url.split("/")[-1].split("?")[0]
            if url_filename and "." in url_filename:
                filename = url_filename
            else:
                filename = f"image_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
        else:
            if not any(filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']):
                filename = f"{filename}.jpg"
        
        file_path = images_dir / filename
        
        # 下载图片
        with httpx.Client(timeout=30.0, follow_redirects=True, proxy=None, trust_env=False) as client:
            response = client.get(url)
            response.raise_for_status()
            
            content_type = response.headers.get("content-type", "")
            if not content_type.startswith("image/"):
                return {"success": False, "error": f"URL 不是图片: {content_type}"}
            
            with open(file_path, "wb") as f:
                f.write(response.content)
        
        return {
            "success": True,
            "message": "图片下载成功",
            "local_path": str(file_path),
            "filename": filename,
            "size": file_path.stat().st_size
        }
        
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"下载失败: HTTP {e.response.status_code}"}
    except Exception as e:
        return {"success": False, "error": f"下载出错: {str(e)}"}


def insert_image(filename: str, image_path: str, width: float = None) -> dict:
    """
    将图片插入到 Word 文档中
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文档不存在: {filename}"}
        
        if not Path(image_path).exists():
            return {"success": False, "error": f"图片不存在: {image_path}"}
        
        doc = Document(str(file_path))
        
        if width:
            doc.add_picture(image_path, width=Inches(width))
        else:
            doc.add_picture(image_path, width=Inches(4))  # 默认宽度 4 英寸
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "图片插入成功",
            "image_path": image_path
        }
    except Exception as e:
        return {"success": False, "error": f"插入图片失败: {str(e)}"}


# ==================== 多 Agent 工具实现 ====================

# 初始化多 Agent Pipeline
_word_tools_for_pipeline = {
    "create_document": create_document,
    "read_document": read_document,
    "update_document": update_document,
    "add_table": add_table,
    "insert_image": insert_image
}
_document_pipeline = DocumentCreationPipeline(
    word_tools=_word_tools_for_pipeline,
    pass_threshold=7,
    max_iterations=3
)


def create_document_with_agents(user_request: str, auto_confirm: bool = False) -> dict:
    """
    使用多 Agent 协作创建文档
    
    流程：
    1. 结构化 Agent：解析用户请求，提取参数
    2. 创作 Agent：生成文档内容
    3. 评审 Agent：评估质量，不达标则重新创作（最多3轮）
    """
    try:
        logger.info(f"[多Agent] 开始处理请求: {user_request[:100]}...")
        
        result = _document_pipeline.run(user_request, auto_confirm=auto_confirm)
        
        # 如果成功且不需要澄清，自动保存文档
        if result.get("success") and result.get("final_draft"):
            draft = result["final_draft"]
            logger.info(f"[多Agent] 创作成功，评分: {result.get('final_review', {}).get('score', 'N/A')}/10")
            
            save_result = create_document(
                filename=draft["filename"],
                title=draft["title"],
                content=draft["content"]
            )
            result["save_result"] = save_result
            
            if save_result.get("success"):
                logger.info(f"[多Agent] 文档已保存: {draft['filename']}")
        
        # 添加 Agent 流程说明
        result["agent_info"] = {
            "pipeline": "结构化Agent → 创作Agent → 评审Agent",
            "pass_threshold": 7,
            "max_iterations": 3
        }
        
        return result
        
    except Exception as e:
        logger.error(f"[多Agent] 出错: {str(e)}")
        return {
            "success": False,
            "error": f"多 Agent 协作出错: {str(e)}"
        }


def structurize_input(user_input: str) -> dict:
    """
    仅使用结构化 Agent 解析用户输入（不创建文档）
    """
    try:
        logger.info(f"[结构化Agent] 解析输入: {user_input[:100]}...")
        
        structurizer = StructurizerAgent()
        task, questions = structurizer.process(user_input)
        
        result = {
            "success": True,
            "task": task.to_dict(),
            "clarification_questions": questions,
            "has_questions": len(questions) > 0
        }
        
        logger.info(f"[结构化Agent] 识别意图: {task.intent}, 问题数: {len(questions)}")
        
        return result
        
    except Exception as e:
        logger.error(f"[结构化Agent] 出错: {str(e)}")
        return {
            "success": False,
            "error": f"结构化解析出错: {str(e)}"
        }


# ==================== 记忆工具实现 ====================

def save_to_memory(content: str, category: str = "fact", importance: float = 0.5, tags: List[str] = None, session_id: str = "default") -> dict:
    """
    保存信息到长期记忆
    """
    try:
        session = get_session(session_id)
        memory_id = session.remember(
            content=content,
            category=category,
            importance=importance,
            tags=tags or []
        )
        
        logger.info(f"[记忆] 保存到长期记忆: {content[:50]}... (ID: {memory_id})")
        
        return {
            "success": True,
            "message": "已保存到长期记忆",
            "memory_id": memory_id,
            "content": content[:100],
            "category": category,
            "importance": importance
        }
    except Exception as e:
        logger.error(f"[记忆] 保存失败: {str(e)}")
        return {
            "success": False,
            "error": f"保存记忆失败: {str(e)}"
        }


def recall_memory(query: str, category: str = None, limit: int = 5, session_id: str = "default") -> dict:
    """
    从长期记忆中检索信息
    """
    try:
        session = get_session(session_id)
        results = session.long_term.search(query, category=category, limit=limit)
        
        memories = []
        for item in results:
            memories.append({
                "id": item.id,
                "content": item.content,
                "category": item.category,
                "importance": item.importance,
                "tags": item.tags,
                "access_count": item.access_count
            })
        
        logger.info(f"[记忆] 检索 '{query}' 找到 {len(memories)} 条记忆")
        
        return {
            "success": True,
            "query": query,
            "count": len(memories),
            "memories": memories
        }
    except Exception as e:
        logger.error(f"[记忆] 检索失败: {str(e)}")
        return {
            "success": False,
            "error": f"检索记忆失败: {str(e)}"
        }


def get_memory_stats(session_id: str = "default") -> dict:
    """
    获取记忆统计信息
    """
    try:
        session = get_session(session_id)
        stats = session.get_stats()
        
        # 添加长期记忆详情
        long_term_summary = session.long_term.get_summary()
        
        return {
            "success": True,
            "session_stats": stats,
            "long_term_summary": long_term_summary,
            "working_memory": session.working.get_summary()
        }
    except Exception as e:
        logger.error(f"[记忆] 获取统计失败: {str(e)}")
        return {
            "success": False,
            "error": f"获取记忆统计失败: {str(e)}"
        }


# 工具映射
TOOL_HANDLERS = {
    "create_document": create_document,
    "read_document": read_document,
    "update_document": update_document,
    "delete_document": delete_document,
    "list_documents": list_documents,
    "add_table": add_table,
    "search_replace": search_replace,
    "google_search": google_search,
    "google_image_search": google_image_search,
    "download_image": download_image,
    "insert_image": insert_image,
    # 多 Agent 工具
    "create_document_with_agents": create_document_with_agents,
    "structurize_input": structurize_input,
    # 记忆工具
    "save_to_memory": save_to_memory,
    "recall_memory": recall_memory,
    "get_memory_stats": get_memory_stats,
}


# ==================== LLM 调用 ====================

def get_tools_for_llm() -> list:
    """将工具定义转换为 OpenAI function calling 格式"""
    tools = []
    for name, info in TOOLS.items():
        tools.append({
            "type": "function",
            "function": {
                "name": name,
                "description": info["description"],
                "parameters": info["parameters"]
            }
        })
    return tools


async def call_llm(messages: list, tools: list = None) -> dict:
    """调用 LLM API"""
    if not LLM_CONFIG:
        raise ValueError("LLM 配置未加载")
    
    base_url = LLM_CONFIG.get("baseURL", "").rstrip("/")
    api_token = LLM_CONFIG.get("apiToken")
    model = LLM_CONFIG.get("model")
    
    if not all([base_url, api_token, model]):
        raise ValueError("LLM 配置不完整")
    
    headers = {
        "Authorization": f"Bearer {api_token}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.7,
    }
    
    if tools:
        payload["tools"] = tools
        payload["tool_choice"] = "auto"
    
    async with httpx.AsyncClient(timeout=60.0, proxy=None, trust_env=False) as client:
        response = await client.post(
            f"{base_url}/chat/completions",
            headers=headers,
            json=payload
        )
        response.raise_for_status()
        return response.json()


def execute_tool(tool_name: str, arguments: dict) -> dict:
    """执行工具调用"""
    if tool_name not in TOOL_HANDLERS:
        return {"success": False, "error": f"未知工具: {tool_name}"}
    
    handler = TOOL_HANDLERS[tool_name]
    return handler(**arguments)


def parse_qwen_tool_calls(reasoning_content: str) -> list:
    """
    解析 Qwen 模型 reasoning_content 中的 <tool_call> 标签
    格式: <tool_call>\n{"name": "xxx", "arguments": {...}}\n</tool_call>
    """
    tool_calls = []
    
    # 匹配 <tool_call>...</tool_call> 中的 JSON
    pattern = r'<tool_call>\s*(\{.*?\})\s*(?:</tool_call>|$)'
    matches = re.findall(pattern, reasoning_content, re.DOTALL)
    
    for idx, json_str in enumerate(matches):
        try:
            # 清理 JSON 字符串（处理可能的截断）
            json_str = json_str.strip()
            
            # 尝试修复不完整的 JSON（如果被截断）
            if not json_str.endswith('}'):
                # 尝试找到最后一个完整的 }
                last_brace = json_str.rfind('}')
                if last_brace > 0:
                    json_str = json_str[:last_brace + 1]
            
            data = json.loads(json_str)
            tool_name = data.get("name")
            arguments = data.get("arguments", {})
            
            if tool_name and tool_name in TOOL_HANDLERS:
                tool_calls.append({
                    "id": f"qwen_call_{idx}",
                    "function": {
                        "name": tool_name,
                        "arguments": json.dumps(arguments, ensure_ascii=False) if isinstance(arguments, dict) else arguments
                    }
                })
        except json.JSONDecodeError as e:
            logger.warning(f"解析工具调用 JSON 失败: {e}, 原始内容: {json_str[:200]}")
            continue
    
    return tool_calls


# ==================== API 端点 ====================

class ToolCallRequest(BaseModel):
    """工具调用请求"""
    tool: str
    params: dict = {}


class AgentRequest(BaseModel):
    """Agent 请求"""
    query: str
    title: Optional[str] = None
    filename: Optional[str] = None
    session_id: Optional[str] = "default"  # 会话 ID，用于记忆隔离


def _sanitize_filename(name: str) -> str:
    """清理文件名"""
    s = re.sub(r"\s+", "_", name.strip())
    s = re.sub(r"[^\w\u4e00-\u9fa5_-]+", "", s)
    return s[:40] if s else ""


@app.get("/")
async def root():
    """服务器状态"""
    return {
        "name": "Word MCP Server",
        "version": "2.0.0",
        "status": "running",
        "llm": {
            "baseURL": LLM_CONFIG.get("baseURL"),
            "model": LLM_CONFIG.get("model")
        },
        "endpoints": {
            "tools": "/tools",
            "call": "/call (POST)",
            "sse": "/sse",
            "sse_agent": "/sse/agent (POST)",
            "documents": "/documents"
        }
    }


@app.get("/tools")
async def get_tools():
    """获取可用工具列表"""
    return {"tools": TOOLS}


@app.post("/call")
async def call_tool(request: ToolCallRequest):
    """直接调用工具（不经过 LLM）"""
    tool_name = request.tool
    params = request.params
    
    if tool_name not in TOOL_HANDLERS:
        return {"success": False, "error": f"未知工具: {tool_name}"}
    
    handler = TOOL_HANDLERS[tool_name]
    result = handler(**params)
    
    logger.info(f"工具调用: {tool_name}({params}) -> {result.get('success')}")
    return result


@app.get("/documents")
async def get_documents():
    """获取文档列表"""
    return list_documents()


@app.get("/sse")
async def sse_endpoint(request: Request):
    """SSE 端点 - 用于实时事件流"""
    
    async def event_generator():
        yield f"data: {json.dumps({'type': 'connected', 'message': 'SSE 连接成功'})}\n\n"
        yield f"data: {json.dumps({'type': 'tools', 'tools': list(TOOLS.keys())})}\n\n"
        
        while True:
            if await request.is_disconnected():
                break
            yield f"data: {json.dumps({'type': 'heartbeat', 'time': datetime.now().isoformat()})}\n\n"
            await asyncio.sleep(30)
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@app.post("/sse/call")
async def sse_call_tool(request: ToolCallRequest):
    """SSE 方式调用工具"""
    
    async def stream_response():
        tool_name = request.tool
        params = request.params
        
        yield f"data: {json.dumps({'type': 'start', 'tool': tool_name})}\n\n"
        
        if tool_name not in TOOL_HANDLERS:
            yield f"data: {json.dumps({'type': 'error', 'error': f'未知工具: {tool_name}'})}\n\n"
            return
        
        handler = TOOL_HANDLERS[tool_name]
        result = handler(**params)
        
        yield f"data: {json.dumps({'type': 'result', 'data': result})}\n\n"
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    
    return StreamingResponse(
        stream_response(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive"
        }
    )


@app.post("/sse/agent")
async def sse_agent(request: AgentRequest, http_request: Request):
    """
    LLM Agent SSE 端点（带记忆功能）
    
    1. 接收用户的自然语言查询
    2. 从记忆系统加载历史上下文
    3. 调用 LLM 理解意图并决定调用哪些工具
    4. 执行工具调用
    5. 将结果返回给 LLM 继续处理
    6. 循环直到 LLM 给出最终答案
    7. 保存对话到记忆系统
    """

    async def stream_response():
        query = request.query
        title = (request.title or "").strip()
        filename_hint = (request.filename or "").strip()
        session_id = request.session_id or "default"
        
        # 获取或创建会话
        session = get_session(session_id)
        logger.info(f"[记忆] 使用会话: {session_id}, 历史消息数: {len(session.short_term)}")
        
        # 构建系统提示 - 强制使用多 Agent 流水线
        system_prompt = """你是一个智能文档助手，使用多 Agent 流水线来创建高质量文档。你具有记忆能力，可以记住之前的对话内容。

## 核心原则：使用多 Agent 流水线

所有文档创建任务必须经过 **三阶段 Agent 流水线**：
1. **结构化 Agent** - 解析用户输入，提取参数，识别缺失信息
2. **创作 Agent** - 根据结构化数据生成文档内容  
3. **评审 Agent** - 评估文档质量（1-10分），不达标则重新创作

## 工具选择规则（必须遵守）

| 任务类型 | 必须使用的工具 |
|----------|----------------|
| 创建新文档 | `create_document_with_agents` ⚠️ 必须使用！ |
| 仅解析用户输入 | `structurize_input` |
| 简单修改已有文档 | `update_document` |
| 读取文档 | `read_document` |
| 删除文档 | `delete_document` |
| 列出文档 | `list_documents` |
| 保存到长期记忆 | `save_to_memory` |
| 从长期记忆检索 | `recall_memory` |

## ⚠️ 禁止行为
❌ **禁止直接使用 `create_document`**，必须用 `create_document_with_agents`
❌ 不要假设文件名、标题、内容
❌ 不要跳过多 Agent 流程

## 记忆功能

你可以：
- 记住用户之前提到的偏好和需求
- 参考之前创建的文档
- 保存重要信息到长期记忆

## 多 Agent 工具使用方法

```python
# 创建文档（必须使用）
create_document_with_agents(
    user_request="用户的原始请求文本",
    auto_confirm=False  # False=需要澄清时会返回问题
)
```

返回值说明：
- `needs_clarification=True`: 需要向用户提问，`questions` 包含问题列表
- `success=True`: 文档创建成功，包含评分和最终内容
- `iterations`: 经过几轮创作-评审循环

## 其他工具

搜索功能：
- google_search(query, num_results): 搜索文字信息
- google_image_search(query, num_results): 搜索图片

图片操作（按顺序使用）：
1. google_image_search → 获取图片URL
2. download_image → 下载到本地
3. insert_image → 插入文档

现在，请使用正确的工具来完成用户的请求。记住：创建文档必须用 `create_document_with_agents`！"""

        # 设置系统提示词到会话
        session.short_term.set_system_prompt(system_prompt)
        
        # 添加用户消息到记忆
        user_message = f"用户请求: {query}" + (f"\n建议标题: {title}" if title else "") + (f"\n建议文件名: {filename_hint}" if filename_hint else "")
        session.add_message("user", user_message)
        
        # 从会话获取完整上下文（包含历史记忆）
        messages = session.get_context_for_llm()
        
        tools = get_tools_for_llm()
        
        yield f"data: {json.dumps({'type': 'start', 'message': '正在理解您的需求...'}, ensure_ascii=False)}\n\n"
        await asyncio.sleep(0)
        
        max_iterations = 10  # 防止无限循环
        iteration = 0
        
        while iteration < max_iterations:
            if await http_request.is_disconnected():
                logger.info("客户端断开连接")
                break
            
            iteration += 1
            
            try:
                # 调用 LLM
                yield f"data: {json.dumps({'type': 'thinking', 'message': f'正在思考 (第 {iteration} 轮)...'}, ensure_ascii=False)}\n\n"
                
                llm_response = await call_llm(messages, tools)
                logger.info(f"LLM 原始响应: {json.dumps(llm_response, ensure_ascii=False)[:500]}")


                choice = llm_response.get("choices", [{}])[0]
                message = choice.get("message", {})
                
                # 检查是否有工具调用
                tool_calls = message.get("tool_calls", [])
                
                # Qwen 模型特殊处理：从 reasoning_content 中解析 <tool_call> 标签
                if not tool_calls:
                    reasoning_content = message.get("reasoning_content", "")
                    if reasoning_content and "<tool_call>" in reasoning_content:
                        tool_calls = parse_qwen_tool_calls(reasoning_content)
                        logger.info(f"从 reasoning_content 解析出工具调用: {tool_calls}")
                
                if tool_calls:
                    # 添加助手消息到历史
                    messages.append(message)
                    
                    # 执行每个工具调用
                    for tool_call in tool_calls:
                        tool_name = tool_call["function"]["name"]
                        tool_id = tool_call["id"]
                        
                        try:
                            arguments = json.loads(tool_call["function"]["arguments"])
                        except json.JSONDecodeError:
                            arguments = {}
                        
                        yield f"data: {json.dumps({'type': 'tool_call', 'tool': tool_name, 'arguments': arguments}, ensure_ascii=False)}\n\n"
                        
                        # 执行工具
                        result = execute_tool(tool_name, arguments)
                        
                        yield f"data: {json.dumps({'type': 'tool_result', 'tool': tool_name, 'result': result}, ensure_ascii=False)}\n\n"
                        
                        # 添加工具结果到消息历史
                        messages.append({
                            "role": "tool",
                            "tool_call_id": tool_id,
                            "content": json.dumps(result, ensure_ascii=False)
                        })
                        
                        await asyncio.sleep(0)
                else:
                    # 没有工具调用，LLM 给出了最终回答
                    content = message.get("content", "")
                    
                    # Qwen 模型特殊处理：如果 content 为空，使用 reasoning_content
                    if not content:
                        reasoning_content = message.get("reasoning_content", "")
                        if reasoning_content and "<tool_call>" not in reasoning_content:
                            content = reasoning_content
                    
                    # 处理 Qwen 模型的思考标签
                    if content:
                        # 移除 <think>...</think> 标签内容
                        content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
                    
                    # 保存助手回复到记忆
                    session.add_message("assistant", content)
                    logger.info(f"[记忆] 保存助手回复到会话 {session_id}")
                    
                    yield f"data: {json.dumps({'type': 'response', 'content': content}, ensure_ascii=False)}\n\n"
                    break
                    
            except httpx.HTTPStatusError as e:
                error_msg = f"LLM API 错误: {e.response.status_code}"
                logger.error(f"{error_msg}: {e.response.text}")
                yield f"data: {json.dumps({'type': 'error', 'error': error_msg}, ensure_ascii=False)}\n\n"
                break
            except Exception as e:
                error_msg = f"处理错误: {str(e)}"
                logger.exception(error_msg)
                yield f"data: {json.dumps({'type': 'error', 'error': error_msg}, ensure_ascii=False)}\n\n"
                break
        
        if iteration >= max_iterations:
            yield f"data: {json.dumps({'type': 'warning', 'message': '达到最大迭代次数限制'}, ensure_ascii=False)}\n\n"
        
        yield f"data: {json.dumps({'type': 'done'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        stream_response(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/chat")
async def chat(request: AgentRequest):
    """
    非流式 Chat 端点（用于简单请求）
    """
    query = request.query
    title = (request.title or "").strip()
    filename_hint = (request.filename or "").strip()
    
    system_prompt = """你是一个专业的 Word 文档助手，使用多 Agent 流水线创建文档。

⚠️ 重要：创建文档必须使用 `create_document_with_agents`，禁止直接用 `create_document`！

多 Agent 流程会自动：
1. 结构化解析用户输入
2. 创作文档内容
3. 评审质量（评分<7自动重写）

使用方法：create_document_with_agents(user_request="用户请求", auto_confirm=False)"""
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"用户请求: {query}" + (f"\n建议标题: {title}" if title else "") + (f"\n建议文件名: {filename_hint}" if filename_hint else "")}
    ]
    
    tools = get_tools_for_llm()
    results = []
    
    max_iterations = 10
    for _ in range(max_iterations):
        try:
            llm_response = await call_llm(messages, tools)
            choice = llm_response.get("choices", [{}])[0]
            message = choice.get("message", {})
            
            tool_calls = message.get("tool_calls", [])
            
            if tool_calls:
                messages.append(message)
                
                for tool_call in tool_calls:
                    tool_name = tool_call["function"]["name"]
                    tool_id = tool_call["id"]
                    
                    try:
                        arguments = json.loads(tool_call["function"]["arguments"])
                    except json.JSONDecodeError:
                        arguments = {}
                    
                    result = execute_tool(tool_name, arguments)
                    results.append({"tool": tool_name, "arguments": arguments, "result": result})
                    
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_id,
                        "content": json.dumps(result, ensure_ascii=False)
                    })
            else:
                content = message.get("content", "")
                if content:
                    content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
                return {
                    "success": True,
                    "response": content,
                    "tool_calls": results
                }
                
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    return {"success": False, "error": "达到最大迭代次数"}


# ==================== 记忆管理 API ====================

@app.get("/memory/sessions")
async def list_sessions():
    """列出所有会话"""
    memory_mgr = MemoryManager()
    return {
        "success": True,
        "sessions": memory_mgr.list_sessions()
    }


@app.get("/memory/session/{session_id}")
async def get_session_info(session_id: str):
    """获取会话信息"""
    session = get_session(session_id)
    return {
        "success": True,
        "stats": session.get_stats(),
        "recent_messages": [msg.to_dict() for msg in session.short_term.get_recent(10)],
        "long_term_summary": session.long_term.get_summary()
    }


@app.delete("/memory/session/{session_id}")
async def delete_session(session_id: str):
    """删除会话"""
    memory_mgr = MemoryManager()
    if memory_mgr.delete_session(session_id):
        return {"success": True, "message": f"会话 {session_id} 已删除"}
    return {"success": False, "error": "会话不存在"}


@app.post("/memory/session/{session_id}/clear")
async def clear_session(session_id: str):
    """清空会话的短期记忆（保留长期记忆）"""
    session = get_session(session_id)
    session.short_term.clear()
    session.working.clear()
    return {"success": True, "message": f"会话 {session_id} 的短期记忆已清空"}


@app.get("/memory/session/{session_id}/history")
async def get_session_history(session_id: str, limit: int = 20):
    """获取会话对话历史"""
    session = get_session(session_id)
    messages = session.short_term.get_recent(limit)
    return {
        "success": True,
        "session_id": session_id,
        "count": len(messages),
        "messages": [msg.to_dict() for msg in messages]
    }


class MemoryAddRequest(BaseModel):
    """添加记忆请求"""
    content: str
    category: str = "fact"
    importance: float = 0.5
    tags: List[str] = []


@app.post("/memory/session/{session_id}/remember")
async def add_memory(session_id: str, request: MemoryAddRequest):
    """手动添加长期记忆"""
    session = get_session(session_id)
    memory_id = session.remember(
        content=request.content,
        category=request.category,
        importance=request.importance,
        tags=request.tags
    )
    return {
        "success": True,
        "memory_id": memory_id,
        "message": "记忆已添加"
    }


@app.get("/memory/session/{session_id}/recall")
async def search_memory(session_id: str, query: str, limit: int = 5):
    """搜索长期记忆"""
    session = get_session(session_id)
    results = session.recall(query, limit=limit)
    return {
        "success": True,
        "query": query,
        "count": len(results),
        "memories": [
            {
                "id": item.id,
                "content": item.content,
                "category": item.category,
                "importance": item.importance,
                "tags": item.tags
            }
            for item in results
        ]
    }


if __name__ == "__main__":
    import uvicorn
    
    print("=" * 50)
    print("Word MCP Server (SSE + LLM Agent)")
    print("=" * 50)
    print(f"文档目录: {WORD_DIR.absolute()}")
    print(f"服务地址: http://localhost:8080")
    print(f"LLM: {LLM_CONFIG.get('model')} @ {LLM_CONFIG.get('baseURL')}")
    print("=" * 50)
    print("\n可用端点:")
    print("  GET  /           - 服务器状态")
    print("  GET  /tools      - 工具列表")
    print("  POST /call       - 直接调用工具")
    print("  GET  /sse        - SSE 连接")
    print("  POST /sse/call   - SSE 调用工具")
    print("  POST /sse/agent  - LLM Agent (SSE 流式, 带记忆)")
    print("  POST /chat       - LLM Agent (非流式)")
    print("  GET  /documents  - 文档列表")
    print("\n记忆管理端点:")
    print("  GET  /memory/sessions              - 列出所有会话")
    print("  GET  /memory/session/{id}          - 获取会话信息")
    print("  DELETE /memory/session/{id}        - 删除会话")
    print("  POST /memory/session/{id}/clear    - 清空短期记忆")
    print("  GET  /memory/session/{id}/history  - 获取对话历史")
    print("  POST /memory/session/{id}/remember - 添加长期记忆")
    print("  GET  /memory/session/{id}/recall   - 搜索长期记忆")
    print("=" * 50)
    
    uvicorn.run(app, host="0.0.0.0", port=8080)
